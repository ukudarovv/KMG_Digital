from pathlib import Path

import fitz  # PyMuPDF
from docx import Document

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


class FileValidationError(ValueError):
    pass


class TextExtractionError(ValueError):
    pass


def validate_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise FileValidationError("Допустимы только файлы .pdf и .docx")
    return ext


def validate_file_size(file_size: int, max_size_bytes: int) -> None:
    if file_size == 0:
        raise FileValidationError("Файл пустой")
    if file_size > max_size_bytes:
        raise FileValidationError(f"Размер файла превышает лимит {max_size_bytes // (1024 * 1024)} MB")


def extract_text_from_pdf(file_path: Path) -> str:
    try:
        with fitz.open(file_path) as doc:
            parts = [page.get_text() for page in doc]
    except Exception as exc:
        raise TextExtractionError(f"Не удалось прочитать PDF: {exc}") from exc

    text = "\n".join(parts).strip()
    if not text:
        raise TextExtractionError("Не удалось извлечь текст из PDF")
    return text


def extract_text_from_docx(file_path: Path) -> str:
    try:
        document = Document(file_path)
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())
    except Exception as exc:
        raise TextExtractionError(f"Не удалось прочитать DOCX: {exc}") from exc

    text = "\n".join(parts).strip()
    if not text:
        raise TextExtractionError("Не удалось извлечь текст из DOCX")
    return text


def extract_text(file_path: Path, extension: str) -> str:
    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    if extension == ".docx":
        return extract_text_from_docx(file_path)
    raise FileValidationError("Неподдерживаемый тип файла")
